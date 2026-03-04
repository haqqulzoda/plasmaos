"""
Plasma AI - Proposals Endpoints

AI-generated proposal management with tier gating.
"""

import asyncio
import io
import logging
from datetime import datetime
from typing import Any
from uuid import UUID

logger = logging.getLogger(__name__)

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from pathlib import Path as _Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

# ── Register Cyrillic-compatible Roboto fonts ──
_FONTS_DIR = _Path(__file__).resolve().parent.parent.parent.parent / "fonts"
pdfmetrics.registerFont(TTFont("Roboto", str(_FONTS_DIR / "Roboto-Regular.ttf")))
pdfmetrics.registerFont(TTFont("Roboto-Bold", str(_FONTS_DIR / "Roboto-Bold.ttf")))
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.api.deps import get_current_user, require_tier
from app.core.evaluator import DynamicComplianceResult
from app.core.security import authenticated_dependency
from app.db.session import get_db
from app.models.all_models import (
    Proposal,
    ProposalStatus,
    RiskOverrideLog,
    SubscriptionTier,
    TaxonomyNode,
    Tender,
    User,
)
from app.models.audit import TenderAnalysis
from app.models.company import CompanyProfile
from app.schemas.proposal import (
    ProposalCreate,
    ProposalResponse,
    ProposalUpdate,
    ProposalWithTenderResponse,
)

router = APIRouter(dependencies=[authenticated_dependency()])


# =============================================================================
# Additional Schemas
# =============================================================================

class AIStrategicLineItem(BaseModel):
    """Single strategic line item in draft response."""
    name: str
    quantity: float = 1
    unit: str = "pcs"
    unit_price: float = 0
    total: float = 0


class AIDraftResponse(BaseModel):
    """Response from AI strategic drafting."""
    strategic_summary: str
    suggested_price: float
    delivery_days: str
    line_items: list[AIStrategicLineItem] = Field(default_factory=list)


class PDFGenerateRequest(BaseModel):
    """Request body for PDF generation."""
    price: float
    delivery_days: int
    company_name: str = "Your Company LLC"


def _analysis_owner_key(
    *,
    current_user: User,
    profile: CompanyProfile | None,
) -> str:
    """
    Build a tenant-safe analysis ownership key.

    TenderAnalysis currently has no explicit user_id column, so we persist and
    query by a deterministic key derived from the authenticated user context.
    """
    profile_token = str(profile.id) if profile is not None else "no-profile"
    return f"{current_user.id}:{profile_token}"


# =============================================================================
# Endpoints
# =============================================================================

@router.post("", response_model=ProposalResponse, status_code=status.HTTP_201_CREATED)
async def create_proposal(
    proposal_data: ProposalCreate,
    current_user: User = Depends(require_tier(SubscriptionTier.SCOUT)),
    db: AsyncSession = Depends(get_db),
) -> ProposalResponse:
    """
    Create a new proposal draft for a tender.
    
    Requires: Agent tier or higher.
    """
    # Verify tender exists
    result = await db.execute(
        select(Tender).where(Tender.id == proposal_data.tender_id)
    )
    tender = result.scalar_one_or_none()
    
    if not tender:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Tender not found",
        )
    
    # Check if user already has a proposal for this tender
    existing_result = await db.execute(
        select(Proposal).where(
            Proposal.user_id == current_user.id,
            Proposal.tender_id == proposal_data.tender_id,
        )
    )
    existing = existing_result.scalar_one_or_none()
    
    if existing:
        # Return existing proposal
        return ProposalResponse.model_validate(existing)
    
    # Create new proposal
    proposal = Proposal(
        user_id=current_user.id,
        tender_id=proposal_data.tender_id,
        status=ProposalStatus.DRAFT,
        ai_confidence_score=0,
        structured_data={},
    )
    db.add(proposal)
    await db.commit()
    await db.refresh(proposal)
    
    return ProposalResponse.model_validate(proposal)


@router.get("", response_model=list[ProposalWithTenderResponse])
async def list_proposals(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> list[ProposalWithTenderResponse]:
    """
    List all proposals for the current user.
    """
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(Proposal.user_id == current_user.id)
        .order_by(Proposal.created_at.desc())
    )
    proposals = result.scalars().all()
    
    response = []
    for p in proposals:
        data = ProposalWithTenderResponse(
            id=p.id,
            user_id=p.user_id,
            tender_id=p.tender_id,
            status=p.status,
            ai_confidence_score=p.ai_confidence_score,
            structured_data=p.structured_data,
            final_pdf_url=p.final_pdf_url,
            margin_percent=p.margin_percent,
            include_vat=p.include_vat,
            currency=p.currency,
            created_at=p.created_at,
            tender_title=p.tender.title if p.tender else "Unknown",
            tender_budget=p.tender.budget if p.tender else 0,
            tender_currency=p.tender.currency if p.tender else "UZS",
            tender_deadline=p.tender.deadline if p.tender else None,
            tender_region=p.tender.region if p.tender else None,
        )
        response.append(data)
    
    return response


@router.get("/{proposal_id}", response_model=ProposalWithTenderResponse)
async def get_proposal(
    proposal_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ProposalWithTenderResponse:
    """
    Get a specific proposal with tender details.
    """
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )
    
    return ProposalWithTenderResponse(
        id=proposal.id,
        user_id=proposal.user_id,
        tender_id=proposal.tender_id,
        status=proposal.status,
        ai_confidence_score=proposal.ai_confidence_score,
        structured_data=proposal.structured_data,
        final_pdf_url=proposal.final_pdf_url,
        margin_percent=proposal.margin_percent,
        include_vat=proposal.include_vat,
        currency=proposal.currency,
        created_at=proposal.created_at,
        tender_title=proposal.tender.title if proposal.tender else "Unknown",
        tender_budget=proposal.tender.budget if proposal.tender else 0,
        tender_currency=proposal.tender.currency if proposal.tender else "UZS",
        tender_deadline=proposal.tender.deadline if proposal.tender else None,
        tender_region=proposal.tender.region if proposal.tender else None,
    )


@router.put("/{proposal_id}", response_model=ProposalResponse)
async def update_proposal(
    proposal_id: UUID,
    update_data: ProposalUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> ProposalResponse:
    """
    Update proposal data with financial calculations.
    
    Calculates:
    - unit_price = base_cost × (1 + margin_percent/100)
    - subtotal = sum of all item totals
    - vat_amount = subtotal × 0.12 (if include_vat)
    - grand_total = subtotal + vat_amount
    """
    result = await db.execute(
        select(Proposal).where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )
    
    # Update financial model fields
    if update_data.margin_percent is not None:
        proposal.margin_percent = update_data.margin_percent
    
    if update_data.include_vat is not None:
        proposal.include_vat = update_data.include_vat
    
    # Get current margin for calculations
    margin = proposal.margin_percent
    include_vat = proposal.include_vat
    
    # Update structured data
    if update_data.structured_data is not None:
        proposal.structured_data = update_data.structured_data
    
    current_data: dict[str, Any] = proposal.structured_data or {}
    
    # Process items with financial calculations
    if update_data.items is not None:
        calculated_items = []
        subtotal = 0.0
        
        for item in update_data.items:
            # Calculate unit sell price from base cost + margin
            unit_price = item.base_cost * (1 + margin / 100)
            total_item = unit_price * item.quantity
            subtotal += total_item
            
            calculated_items.append({
                "name": item.name,
                "unit": item.unit,
                "quantity": item.quantity,
                "base_cost": item.base_cost,
                "unit_price": round(unit_price, 2),
                "total": round(total_item, 2),
            })
        
        # Calculate VAT and grand total
        vat_amount = subtotal * 0.12 if include_vat else 0.0
        grand_total = subtotal + vat_amount
        
        # Store calculated values
        current_data["ai_items"] = calculated_items
        current_data["subtotal"] = round(subtotal, 2)
        current_data["vat_amount"] = round(vat_amount, 2)
        current_data["grand_total"] = round(grand_total, 2)
        current_data["our_price"] = round(grand_total, 2)
    
    # Update individual fields
    if update_data.our_price is not None:
        current_data["our_price"] = update_data.our_price
    
    if update_data.delivery_days is not None:
        current_data["delivery_days"] = update_data.delivery_days
    
    proposal.structured_data = current_data
    
    await db.commit()
    await db.refresh(proposal)
    
    return ProposalResponse.model_validate(proposal)


@router.post("/{proposal_id}/ai-draft", response_model=AIDraftResponse)
async def ai_draft_proposal(
    proposal_id: UUID,
    force: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> AIDraftResponse:
    """
    Generate a strategic proposal draft using compliance and liability ledger context.
    """
    from app.core.ai import draft_strategic_proposal_async

    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()

    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )

    if not proposal.tender:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal has no associated tender",
        )

    current_data: dict[str, Any] = proposal.structured_data or {}
    if (
        not force
        and isinstance(current_data.get("strategic_summary"), str)
        and current_data.get("line_items")
    ):
        return AIDraftResponse(
            strategic_summary=current_data.get("strategic_summary", ""),
            suggested_price=float(current_data.get("our_price", 0.0)),
            delivery_days=str(current_data.get("delivery_days", "30 calendar days")),
            line_items=[
                AIStrategicLineItem(
                    name=str(item.get("name", "Line Item")),
                    quantity=float(item.get("quantity", 1)),
                    unit=str(item.get("unit", "lot")),
                    unit_price=float(item.get("unit_price", 0)),
                    total=float(item.get("total", 0)),
                )
                for item in current_data.get("line_items", [])
                if isinstance(item, dict)
            ],
        )

    tender_text = (proposal.tender.compiled_master_text or "").strip()
    if not tender_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Tender documents have not been synchronized. Please sync documents first.",
        )

    profile_result = await db.execute(
        select(CompanyProfile).where(CompanyProfile.user_id == current_user.id)
    )
    profile = profile_result.scalar_one_or_none()
    analysis_owner_key = _analysis_owner_key(
        current_user=current_user,
        profile=profile,
    )
    company_name = str(
        (profile.company_name if profile else None)
        or current_user.company_name
        or current_user.name
        or "Unknown Company"
    )

    analysis_result = await db.execute(
        select(TenderAnalysis)
        .where(
            TenderAnalysis.tender_id == proposal.tender_id,
            TenderAnalysis.company_name == analysis_owner_key,
        )
        .order_by(TenderAnalysis.created_at.desc())
        .limit(1)
    )
    latest_analysis = analysis_result.scalar_one_or_none()
    evaluation_payload = (
        (latest_analysis.analysis_json or {}).get("evaluation", {})
        if latest_analysis
        else {}
    )
    try:
        compliance_result = DynamicComplianceResult.model_validate(evaluation_payload)
    except Exception:
        compliance_result = DynamicComplianceResult(
            is_compliant=False,
            met_requirements=[],
            missing_requirements=[],
            unmapped_requirements=[],
            status_message="No cached compliance analysis found.",
        )

    override_result = await db.execute(
        select(RiskOverrideLog, TaxonomyNode)
        .join(TaxonomyNode, RiskOverrideLog.missing_node_id == TaxonomyNode.id)
        .where(
            RiskOverrideLog.tender_id == proposal.tender_id,
            RiskOverrideLog.user_id == current_user.id,
        )
        .order_by(RiskOverrideLog.created_at.asc())
    )
    override_rows = override_result.all()
    accepted_liabilities: list[str] = []
    for override_log, node in override_rows:
        item = f"{node.name} (node_id={override_log.missing_node_id})"
        if override_log.justification:
            item = f"{item}; justification={override_log.justification}"
        accepted_liabilities.append(item)

    compliance_ledger = {
        "evaluation": compliance_result.model_dump(mode="json"),
        "accepted_liabilities": accepted_liabilities,
    }
    company_context = {
        "company_name": company_name,
        "core_services": getattr(current_user, "core_services", "") or "",
        "past_experience": getattr(current_user, "past_experience", "") or "",
    }
    ai_result = await draft_strategic_proposal_async(
        tender_text,
        company_context=company_context,
        compliance_ledger=compliance_ledger,
        accepted_liabilities=accepted_liabilities,
        tender_budget=proposal.tender.budget,
    )

    strategic_summary = str(ai_result.get("strategic_summary", "")).strip()
    if not strategic_summary:
        strategic_summary = (
            "Our team can execute this opportunity with disciplined delivery, "
            "verified credentials, and transparent commercial controls."
        )

    try:
        suggested_price = float(ai_result.get("suggested_price", proposal.tender.budget * 0.85))
    except (TypeError, ValueError):
        suggested_price = float(proposal.tender.budget * 0.85)

    delivery_days = str(ai_result.get("delivery_days", "")).strip() or "30 calendar days"

    raw_line_items = ai_result.get("line_items", [])
    normalized_items: list[dict[str, Any]] = []
    if isinstance(raw_line_items, list):
        for raw_item in raw_line_items:
            if not isinstance(raw_item, dict):
                continue
            try:
                quantity = float(raw_item.get("quantity", 1))
            except (TypeError, ValueError):
                quantity = 1.0
            try:
                unit_price = float(raw_item.get("unit_price", 0))
            except (TypeError, ValueError):
                unit_price = 0.0
            try:
                total = float(raw_item.get("total", quantity * unit_price))
            except (TypeError, ValueError):
                total = quantity * unit_price
            normalized_items.append(
                {
                    "name": str(raw_item.get("name", "Line Item")).strip() or "Line Item",
                    "quantity": quantity,
                    "unit": str(raw_item.get("unit", "lot")).strip() or "lot",
                    "unit_price": unit_price,
                    "total": total,
                }
            )

    if not normalized_items:
        normalized_items = [
            {
                "name": "Delivery Scope",
                "quantity": 1.0,
                "unit": "lot",
                "unit_price": round(suggested_price, 2),
                "total": round(suggested_price, 2),
            }
        ]

    confidence = 90 if compliance_result.is_compliant else 82
    if "error" in ai_result:
        confidence = 60

    current_data["strategic_summary"] = strategic_summary
    current_data["ai_summary"] = strategic_summary
    current_data["our_price"] = suggested_price
    current_data["delivery_days"] = delivery_days
    current_data["line_items"] = normalized_items
    current_data["ai_items"] = normalized_items
    current_data["compliance_ledger"] = compliance_ledger
    current_data["accepted_liabilities"] = accepted_liabilities
    proposal.structured_data = current_data
    proposal.ai_confidence_score = confidence
    await db.commit()

    return AIDraftResponse(
        strategic_summary=strategic_summary,
        suggested_price=suggested_price,
        delivery_days=delivery_days,
        line_items=[
            AIStrategicLineItem(
                name=item["name"],
                quantity=float(item["quantity"]),
                unit=item["unit"],
                unit_price=float(item["unit_price"]),
                total=float(item["total"]),
            )
            for item in normalized_items
        ],
    )

    # ── Cache check: return existing AI results if present ──
    # ── Read pre-parsed text from DB (populated by Celery worker) ──
    # ── Analyze with Gemini via text (no Playwright, no file upload) ──
@router.post("/{proposal_id}/upload-tz", response_model=AIDraftResponse)
async def upload_tender_tz(
    proposal_id: UUID,
    file: UploadFile,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> AIDraftResponse:
    """
    Upload a Technical Task PDF for tenders with archive (ZIP/RAR) documents.
    
    1. Saves uploaded PDF to backend/uploads/{user_id}/{proposal_id}.pdf
    2. Stores extracted text inside the proposal's structured_data
    3. Runs AI analysis on the uploaded file
    4. Returns AI analysis result (same as ai-draft endpoint)
    """
    from pathlib import Path
    
    from app.core.parser import extract_text_from_file
    
    # Fetch proposal with tender
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )
    
    if not proposal.tender:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal has no associated tender",
        )
    
    # Validate file type
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF files are accepted",
        )
    
    # Save file to tenant-scoped uploads directory
    uploads_dir = Path(__file__).parent.parent.parent.parent / "uploads"
    tenant_uploads_dir = uploads_dir / str(current_user.id)
    tenant_uploads_dir.mkdir(parents=True, exist_ok=True)
    file_path = tenant_uploads_dir / f"{proposal.id}.pdf"
    
    # Write file
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    current_data: dict[str, Any] = proposal.structured_data or {}
    current_data["uploaded_tz_path"] = str(file_path)

    # Extract text and persist under this proposal only.
    try:
        extracted_text = await extract_text_from_file(str(file_path))
        if extracted_text and extracted_text.strip():
            current_data["uploaded_tz_text"] = extracted_text.strip()
            logger.info(f"[UPLOAD-TZ] Stored tenant-scoped text ({len(extracted_text)} chars)")
    except Exception as parse_exc:
        logger.warning(f"[UPLOAD-TZ] Text extraction failed: {parse_exc}")
    
    # Build company context from user profile
    company_context = {
        "company_name": current_user.company_name or "",
        "core_services": getattr(current_user, 'core_services', '') or "",
        "past_experience": getattr(current_user, 'past_experience', '') or "",
    }
    
    # Analyze with Gemini AI (direct file upload - handles scanned PDFs!)
    print(f"[UPLOAD-TZ] Analyzing file with Gemini: {file_path}")
    from app.core.ai import analyze_tender_file_async
    ai_result = await analyze_tender_file_async(str(file_path), company_context)
    
    # Calculate estimates based on AI analysis
    tender_budget = proposal.tender.budget
    items = ai_result.get("items", [])
    delivery_days = ai_result.get("delivery_days", 30)
    
    # Estimate cost (75% of budget as baseline)
    estimated_cost = tender_budget * 0.75
    suggested_price = tender_budget * 0.85  # 15% margin
    
    # Build technical summary
    summary = ai_result.get("summary", "Analysis complete.")
    requirements = ai_result.get("required_licenses", [])
    risks = ai_result.get("risks", [])
    
    if requirements:
        summary += f" Required: {', '.join(requirements[:3])}."
    if risks:
        summary += f" Risks: {risks[0]}."
    
    # Determine confidence score
    confidence = 85
    if "error" in ai_result:
        confidence = 50
    elif len(items) == 0:
        confidence = 65
    elif len(items) > 5:
        confidence = 90
    
    strategic_summary = (
        f"{summary[:420]} "
        "This recommendation is backed by verified credentials and practical scope control."
    ).strip()
    delivery_days_text = f"{delivery_days} calendar days"

    item_count = max(len(items), 1)
    unit_price = suggested_price / item_count if item_count else suggested_price
    line_items = [
        {
            "name": str(item.get("name", "Line Item")),
            "quantity": float(item.get("quantity", 1)),
            "unit": str(item.get("unit", "lot")),
            "unit_price": round(unit_price, 2),
            "total": round(unit_price * float(item.get("quantity", 1)), 2),
        }
        for item in items
    ]
    if not line_items:
        line_items = [
            {
                "name": "Delivery Scope",
                "quantity": 1.0,
                "unit": "lot",
                "unit_price": round(suggested_price, 2),
                "total": round(suggested_price, 2),
            }
        ]

    current_data["strategic_summary"] = strategic_summary
    current_data["our_price"] = suggested_price
    current_data["delivery_days"] = delivery_days_text
    current_data["line_items"] = line_items
    current_data["ai_items"] = line_items
    proposal.structured_data = current_data
    proposal.ai_confidence_score = confidence

    await db.commit()

    return AIDraftResponse(
        strategic_summary=strategic_summary,
        suggested_price=suggested_price,
        delivery_days=delivery_days_text,
        line_items=[
            AIStrategicLineItem(
                name=item["name"],
                quantity=float(item["quantity"]),
                unit=item["unit"],
                unit_price=float(item["unit_price"]),
                total=float(item["total"]),
            )
            for item in line_items
        ],
    )


@router.get("/{proposal_id}/uploaded-tz")
async def get_uploaded_tz(
    proposal_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Serve the uploaded Technical Task PDF for preview.
    
    Returns the PDF file that was uploaded via /upload-tz endpoint.
    """
    from pathlib import Path
    from fastapi.responses import Response
    
    # Fetch proposal with tender
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )
    
    if not proposal.tender:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Proposal has no associated tender",
        )
    
    # Look for uploaded PDF
    uploads_dir = Path(__file__).parent.parent.parent.parent / "uploads"
    tenant_uploads_dir = uploads_dir / str(current_user.id)
    file_path = tenant_uploads_dir / f"{proposal.id}.pdf"
    
    if not file_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No uploaded PDF found for this tender",
        )
    
    # Read file and return as Response (same pattern as working document download)
    file_bytes = file_path.read_bytes()
    
    return Response(
        content=file_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename=tz_{proposal.id}.pdf"},
    )


@router.post("/{proposal_id}/generate-pdf")
async def generate_proposal_pdf(
    proposal_id: UUID,
    pdf_data: PDFGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """
    Generate a professional Commercial Proposal (KP) PDF.
    
    Uses company details from user profile and AI-extracted items from proposal.
    Returns a downloadable PDF document suitable for tender submissions.
    """
    # Fetch proposal with tender
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()
    
    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )
    
    tender = proposal.tender
    if not tender:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Tender not found for this proposal",
        )
    
    # Get company details from user profile (fallback to pdf_data for backwards compat)
    company_name = current_user.company_name or pdf_data.company_name
    director_name = current_user.director_name or "Director"
    company_address = current_user.address or ""
    
    # Get AI-extracted items from proposal structured_data
    structured_data = proposal.structured_data or {}
    ai_items = structured_data.get("ai_items", [])
    our_price = structured_data.get("our_price", pdf_data.price)
    delivery_days = structured_data.get("delivery_days", pdf_data.delivery_days)
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        topMargin=1.5*cm, 
        bottomMargin=1.5*cm,
        leftMargin=2*cm,
        rightMargin=2*cm,
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles (Roboto for Cyrillic / Latin-Extended support)
    title_style = ParagraphStyle(
        'KPTitle',
        parent=styles['Heading1'],
        fontName='Roboto-Bold',
        fontSize=16,
        spaceAfter=20,
        alignment=1,  # Center
        textColor=colors.HexColor('#1a1a1a'),
    )
    header_style = ParagraphStyle(
        'HeaderBold',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Roboto-Bold',
        textColor=colors.HexColor('#1a1a1a'),
    )
    normal_style = ParagraphStyle(
        'KPNormal',
        parent=styles['Normal'],
        fontName='Roboto',
        fontSize=10,
        textColor=colors.HexColor('#333333'),
        spaceAfter=6,
    )
    small_style = ParagraphStyle(
        'KPSmall',
        parent=styles['Normal'],
        fontName='Roboto',
        fontSize=9,
        textColor=colors.HexColor('#666666'),
    )
    
    elements = []
    
    # ========== HEADER ==========
    # Company name (left) and Date (right) on same line using table
    header_data = [[
        Paragraph(f"<b>{company_name}</b>", header_style),
        Paragraph(f"<b>Sana:</b> {datetime.now().strftime('%d.%m.%Y')}", normal_style),
    ]]
    header_table = Table(header_data, colWidths=[10*cm, 6*cm])
    header_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    elements.append(header_table)
    
    # Company address if available
    if company_address:
        elements.append(Paragraph(company_address, small_style))
    elements.append(Spacer(1, 20))
    
    # ========== TITLE ==========
    elements.append(Paragraph(
        f"TIJORAT TAKLIFI (COMMERCIAL PROPOSAL)<br/>№ KP-{str(proposal_id)[:8].upper()}",
        title_style
    ))
    elements.append(Spacer(1, 15))
    
    # ========== INTRO ==========
    intro_text = f"Biz, <b>{company_name}</b>, {tender.external_id} raqamli tender bo'yicha quyidagi xizmatlarni taklif etamiz:"
    elements.append(Paragraph(intro_text, normal_style))
    elements.append(Spacer(1, 15))
    
    # ========== STRATEGIC SUMMARY ==========
    strategic_summary_text = (structured_data.get("strategic_summary") or "").strip()
    if strategic_summary_text:
        summary_style = ParagraphStyle(
            'KPSummary',
            parent=normal_style,
            fontName='Roboto',
            fontSize=10,
            leading=15,
            spaceBefore=0,
            spaceAfter=8,
            textColor=colors.HexColor('#222222'),
        )
        elements.append(Paragraph("<b>Strategik tavsif (Executive Summary):</b>", header_style))
        elements.append(Spacer(1, 8))

        # Split wall-of-text into proper paragraphs (~2-3 sentences each)
        import re
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', strategic_summary_text) if s.strip()]
        paragraphs: list[str] = []
        current: list[str] = []
        for sentence in sentences:
            current.append(sentence)
            if len(current) >= 3:
                paragraphs.append(" ".join(current))
                current = []
        if current:
            paragraphs.append(" ".join(current))

        # Render each paragraph inside a left-accent border block
        summary_cells = []
        for para_text in paragraphs:
            summary_cells.append([Paragraph(para_text, summary_style)])

        summary_block = Table(summary_cells, colWidths=[15.5 * cm])
        summary_block.setStyle(TableStyle([
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LINEBEFOREPADDDING', (0, 0), (0, -1), 0),
            ('LINEBEFORE', (0, 0), (0, -1), 2.5, colors.HexColor('#4F46E5')),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8F8FF')),
        ]))
        elements.append(summary_block)
        elements.append(Spacer(1, 18))
    
    # ========== ITEMS TABLE ==========
    if ai_items and len(ai_items) > 0:
        # Calculate unit prices (distribute total evenly if no individual prices)
        total_qty = sum(item.get("quantity", 1) for item in ai_items)
        
        # Table headers
        items_table_data = [
            ["#", "Nomi (Name)", "Birlik", "Miqdor", "Narxi", "Jami"],
        ]
        
        running_total = 0
        for idx, item in enumerate(ai_items, 1):
            name = item.get("name", "Item")[:50]  # Truncate long names
            unit = item.get("unit", "dona")
            qty = item.get("quantity", 1)
            # Use pre-calculated unit_price if available, else estimate
            unit_price = item.get("unit_price", (our_price * qty / total_qty) / qty if total_qty > 0 else our_price / len(ai_items))
            item_total = item.get("total", unit_price * qty)
            running_total += item_total
            
            items_table_data.append([
                str(idx),
                name,
                unit,
                f"{qty:,.0f}",
                f"{unit_price:,.0f}",
                f"{item_total:,.0f}",
            ])
        
        # Get calculated values from structured_data
        subtotal = structured_data.get("subtotal", running_total)
        vat_amount = structured_data.get("vat_amount", 0)
        grand_total = structured_data.get("grand_total", our_price)
        
        # Subtotal row
        items_table_data.append([
            "", "", "", "", Paragraph("<b>Jami summa:</b>", normal_style), 
            Paragraph(f"<b>{subtotal:,.0f}</b>", normal_style)
        ])
        
        # VAT row (only if VAT is included)
        if proposal.include_vat and vat_amount > 0:
            items_table_data.append([
                "", "", "", "", Paragraph("QQS (12%):", normal_style), 
                Paragraph(f"+{vat_amount:,.0f}", normal_style)
            ])
        
        # Grand Total row
        items_table_data.append([
            "", "", "", "", Paragraph("<b>YAKUNIY JAMI:</b>", normal_style), 
            Paragraph(f"<b>{grand_total:,.0f} {tender.currency}</b>", normal_style)
        ])
        
        items_table = Table(
            items_table_data, 
            colWidths=[1*cm, 7*cm, 2*cm, 2*cm, 2.5*cm, 2.5*cm]
        )
        items_table.setStyle(TableStyle([
            # Header row
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F46E5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Roboto-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            # Data rows
            ('BACKGROUND', (0, 1), (-1, -2), colors.HexColor('#FAFAFA')),
            ('FONTNAME', (0, 1), (-1, -1), 'Roboto'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # # column
            ('ALIGN', (2, 1), (-1, -1), 'CENTER'),  # Numeric columns
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            # Total row
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#E8E8FF')),
            ('FONTNAME', (0, -1), (-1, -1), 'Roboto-Bold'),
            # Grid
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
        ]))
        elements.append(items_table)
    else:
        # Fallback: Simple summary table if no items
        summary_data = [
            ["Tavsif (Description)", "Qiymat (Value)"],
            ["Tender byudjeti", f"{tender.budget:,.0f} {tender.currency}"],
            ["Bizning narximiz", f"{our_price:,.0f} {tender.currency}"],
        ]
        summary_table = Table(summary_data, colWidths=[8*cm, 8*cm])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F46E5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Roboto-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FAFAFA')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(summary_table)
    
    elements.append(Spacer(1, 20))
    
    # ========== DELIVERY & TERMS ==========
    elements.append(Paragraph("<b>Yetkazib berish muddati:</b>", normal_style))
    elements.append(Paragraph(f"{delivery_days} ish kuni", normal_style))
    elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>To'lov shartlari:</b>", normal_style))
    elements.append(Paragraph("15% oldindan to'lov, 85% yetkazib berilgandan keyin", normal_style))
    elements.append(Spacer(1, 10))
    
    # VAT notice if applicable
    if proposal.include_vat:
        elements.append(Paragraph("<b>QQS:</b> Narxlar 12% QQSni o'z ichiga oladi (Prices include 12% VAT)", normal_style))
        elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>Taklif amal qilish muddati:</b> 30 kun", normal_style))
    elements.append(Spacer(1, 30))
    
    # ========== SIGNATURE ==========
    elements.append(Paragraph("<b>Direktor:</b>", normal_style))
    elements.append(Spacer(1, 20))
    
    sig_data = [[
        Paragraph(f"{director_name}", normal_style),
        "_" * 25,
        "(imzo / signature)"
    ]]
    sig_table = Table(sig_data, colWidths=[6*cm, 5*cm, 5*cm])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
    ]))
    elements.append(sig_table)
    
    elements.append(Spacer(1, 20))
    
    # Stamp placeholder
    elements.append(Paragraph("M.O. (Muhr joyi / Stamp)", small_style))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    
    # Update proposal status
    proposal.status = ProposalStatus.COMPLETED
    await db.commit()
    
    filename = f"KP_{tender.external_id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
        },
    )


@router.post("/{proposal_id}/export/docx")
async def export_proposal_docx(
    proposal_id: UUID,
    pdf_data: PDFGenerateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> StreamingResponse:
    """
    Export a Commercial Proposal as a Word (.docx) document.

    Mirrors the PDF layout: header, strategic summary paragraphs,
    line-items table, delivery/payment terms, and signature block.
    """
    import re

    from docx import Document as DocxDocument
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    # ── Fetch proposal ──
    result = await db.execute(
        select(Proposal)
        .options(selectinload(Proposal.tender))
        .where(
            Proposal.id == proposal_id,
            Proposal.user_id == current_user.id,
        )
    )
    proposal = result.scalar_one_or_none()

    if not proposal:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Proposal not found",
        )

    tender = proposal.tender
    if not tender:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Tender not found for this proposal",
        )

    company_name = current_user.company_name or pdf_data.company_name
    director_name = current_user.director_name or "Director"
    company_address = current_user.address or ""

    structured_data = proposal.structured_data or {}
    ai_items = structured_data.get("ai_items", [])
    our_price = structured_data.get("our_price", pdf_data.price)
    delivery_days = structured_data.get("delivery_days", pdf_data.delivery_days)

    # ── Build DOCX ──
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    def _style_run(run, size=10, bold=False, color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Roboto"
        if color:
            run.font.color.rgb = RGBColor(*color)

    # ========== HEADER ==========
    header_para = doc.add_paragraph()
    run_company = header_para.add_run(company_name)
    _style_run(run_company, size=12, bold=True)
    header_para.add_run("    ")
    run_date = header_para.add_run(
        f"Sana: {datetime.now().strftime('%d.%m.%Y')}"
    )
    _style_run(run_date, size=10)
    header_para.paragraph_format.space_after = Pt(4)

    if company_address:
        addr_para = doc.add_paragraph()
        run_addr = addr_para.add_run(company_address)
        _style_run(run_addr, size=9, color=(102, 102, 102))
        addr_para.paragraph_format.space_after = Pt(12)

    # ========== TITLE ==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.add_run(
        "TIJORAT TAKLIFI (COMMERCIAL PROPOSAL)"
        f"\n\u2116 KP-{str(proposal_id)[:8].upper()}"
    )
    _style_run(run_title, size=16, bold=True, color=(26, 26, 26))
    title_para.paragraph_format.space_after = Pt(16)

    # ========== INTRO ==========
    intro_para = doc.add_paragraph()
    intro_para.add_run("Biz, ")
    run_cn = intro_para.add_run(company_name)
    _style_run(run_cn, size=10, bold=True)
    rest_run = intro_para.add_run(
        f", {tender.external_id} raqamli tender "
        "bo\u2018yicha quyidagi xizmatlarni taklif etamiz:"
    )
    _style_run(rest_run, size=10)
    intro_para.paragraph_format.space_after = Pt(12)

    # ========== STRATEGIC SUMMARY ==========
    strategic_text = (
        structured_data.get("strategic_summary") or ""
    ).strip()
    if strategic_text:
        heading_para = doc.add_paragraph()
        run_h = heading_para.add_run(
            "Strategik tavsif (Executive Summary):"
        )
        _style_run(run_h, size=12, bold=True, color=(26, 26, 26))
        heading_para.paragraph_format.space_after = Pt(6)

        sentences = [
            s.strip()
            for s in re.split(r'(?<=[.!?])\s+', strategic_text)
            if s.strip()
        ]
        chunks: list[str] = []
        buf: list[str] = []
        for s in sentences:
            buf.append(s)
            if len(buf) >= 3:
                chunks.append(" ".join(buf))
                buf = []
        if buf:
            chunks.append(" ".join(buf))

        for chunk in chunks:
            p = doc.add_paragraph()
            r = p.add_run(chunk)
            _style_run(r, size=10, color=(34, 34, 34))
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Cm(0.5)

    # ========== LINE ITEMS TABLE ==========
    if ai_items and len(ai_items) > 0:
        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        headers = [
            "#", "Nomi (Name)", "Birlik",
            "Miqdor", "Narxi", "Jami",
        ]
        hdr_cells = table.rows[0].cells
        for i, txt in enumerate(headers):
            hdr_cells[i].text = txt
            for para in hdr_cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in para.runs:
                    _style_run(r, size=9, bold=True, color=(255, 255, 255))
            tc_pr = hdr_cells[i]._element.get_or_add_tcPr()
            shd = tc_pr.makeelement(
                qn("w:shd"),
                {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "4F46E5",
                },
            )
            tc_pr.append(shd)

        running_total = 0.0
        for idx, item in enumerate(ai_items, 1):
            name = str(item.get("name", "Item"))[:50]
            unit = item.get("unit", "dona")
            qty = item.get("quantity", 1)
            up = item.get("unit_price", 0)
            itot = item.get("total", up * qty)
            running_total += itot

            row = table.add_row()
            vals = [
                str(idx), name, str(unit),
                f"{qty:,.0f}", f"{up:,.0f}", f"{itot:,.0f}",
            ]
            for i, v in enumerate(vals):
                row.cells[i].text = v
                for para in row.cells[i].paragraphs:
                    if i >= 3:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for r in para.runs:
                        _style_run(r, size=9)

        subtotal = structured_data.get("subtotal", running_total)
        grand_total = structured_data.get("grand_total", our_price)

        row_s = table.add_row()
        row_s.cells[4].text = "Jami summa:"
        row_s.cells[5].text = f"{subtotal:,.0f}"
        for para in row_s.cells[4].paragraphs:
            for r in para.runs:
                _style_run(r, size=9, bold=True)
        for para in row_s.cells[5].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in para.runs:
                _style_run(r, size=9, bold=True)

        row_g = table.add_row()
        row_g.cells[4].text = "YAKUNIY JAMI:"
        row_g.cells[5].text = f"{grand_total:,.0f} {tender.currency}"
        for para in row_g.cells[4].paragraphs:
            for r in para.runs:
                _style_run(r, size=9, bold=True)
        for para in row_g.cells[5].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in para.runs:
                _style_run(r, size=9, bold=True)

        doc.add_paragraph()

    # ========== DELIVERY & TERMS ==========
    p_del = doc.add_paragraph()
    rl = p_del.add_run("Yetkazib berish muddati: ")
    _style_run(rl, size=10, bold=True)
    rv = p_del.add_run(f"{delivery_days} ish kuni")
    _style_run(rv, size=10)

    p_pay = doc.add_paragraph()
    rl2 = p_pay.add_run("To\u2018lov shartlari: ")
    _style_run(rl2, size=10, bold=True)
    rv2 = p_pay.add_run(
        "15% oldindan to\u2018lov, 85% yetkazib berilgandan keyin"
    )
    _style_run(rv2, size=10)

    if proposal.include_vat:
        p_vat = doc.add_paragraph()
        rl3 = p_vat.add_run("QQS: ")
        _style_run(rl3, size=10, bold=True)
        rv3 = p_vat.add_run(
            "Narxlar 12% QQSni o\u2018z ichiga oladi "
            "(Prices include 12% VAT)"
        )
        _style_run(rv3, size=10)

    p_val = doc.add_paragraph()
    rl4 = p_val.add_run("Taklif amal qilish muddati: ")
    _style_run(rl4, size=10, bold=True)
    rv4 = p_val.add_run("30 kun")
    _style_run(rv4, size=10)
    p_val.paragraph_format.space_after = Pt(24)

    # ========== SIGNATURE ==========
    p_dir = doc.add_paragraph()
    rd = p_dir.add_run("Direktor:")
    _style_run(rd, size=10, bold=True)
    p_dir.paragraph_format.space_after = Pt(20)

    sig_para = doc.add_paragraph()
    sr = sig_para.add_run(
        f"{director_name}          "
        f"{'_' * 25}          (imzo / signature)"
    )
    _style_run(sr, size=10)
    sig_para.paragraph_format.space_after = Pt(16)

    stamp_para = doc.add_paragraph()
    rs = stamp_para.add_run("M.O. (Muhr joyi / Stamp)")
    _style_run(rs, size=9, color=(102, 102, 102))

    # ── Write to buffer ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    proposal.status = ProposalStatus.COMPLETED
    await db.commit()

    filename = (
        f"KP_{tender.external_id}"
        f"_{datetime.now().strftime('%Y%m%d')}.docx"
    )

    return StreamingResponse(
        buffer,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f"attachment; filename={filename}",
        },
    )
