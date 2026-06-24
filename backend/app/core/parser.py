"""
Phase 1 parser: unpack archives and parse tender technical documents.
"""

from __future__ import annotations

import asyncio
import io
import logging
import os
import re
import shutil
import tempfile
import time
from pathlib import Path
from typing import BinaryIO, TypedDict
from urllib.parse import urlparse
from zipfile import BadZipFile, ZipFile

import docx
import fitz
import httpx
import pkgutil
from google import genai

# --- Python 3.14 compatibility shim ---
# pkgutil.find_loader was removed in Python 3.14. pytesseract 0.3.10 still
# relies on it, so we restore it here before importing pytesseract.
if not hasattr(pkgutil, "find_loader"):
    import importlib.util

    def _find_loader_shim(fullname: str):  # type: ignore[return]
        spec = importlib.util.find_spec(fullname)
        return spec.loader if spec is not None else None

    pkgutil.find_loader = _find_loader_shim  # type: ignore[attr-defined]

import pytesseract
import rarfile
from pdf2image import convert_from_bytes
from PIL import Image

logger = logging.getLogger(__name__)


def _env_flag(name: str) -> bool:
    return os.environ.get(name, "").strip().lower() in {"1", "true", "yes", "on"}


def _env_int(
    name: str,
    default: int,
    *,
    min_value: int | None = None,
    max_value: int | None = None,
) -> int:
    raw_value = os.environ.get(name)
    if raw_value is None or not raw_value.strip():
        return default

    try:
        value = int(raw_value)
    except ValueError:
        logger.warning("Invalid integer for %s=%r; using default %s", name, raw_value, default)
        return default

    if min_value is not None:
        value = max(min_value, value)
    if max_value is not None:
        value = min(max_value, value)
    return value

SUPPORTED_DOC_SUFFIXES = {".pdf", ".docx"}
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".rtf"}
ARCHIVE_SUFFIXES = {".zip", ".rar"}
ARCHIVE_MIME_TYPES = {
    "application/zip",
    "application/x-zip-compressed",
    "application/vnd.rar",
    "application/x-rar",
    "application/x-rar-compressed",
}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
PDF_MIME_TYPES = {"application/pdf"}
OCR_SAFE_SUFFIXES = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}
OCR_MIN_TEXT_LEN = 50
OCR_LANGS = "uzb+rus+eng"
OCR_PAGE_TIMEOUT_SECONDS = _env_int("TENDER_OCR_PAGE_TIMEOUT_SECONDS", 12, min_value=1, max_value=60)
OCR_MAX_PAGES = _env_int("TENDER_OCR_MAX_PAGES", 2, min_value=0, max_value=25)
OCR_RENDER_DPI = _env_int("TENDER_OCR_RENDER_DPI", 150, min_value=100, max_value=300)
OCR_SKIP_AFTER_TEXT_CHARS = _env_int(
    "TENDER_OCR_SKIP_AFTER_TEXT_CHARS",
    5000,
    min_value=0,
    max_value=1_000_000,
)
OCR_ENABLED = not _env_flag("TENDER_OCR_DISABLED")
DEMO_OCR_BYPASS_ENV = "DEMO_OCR_BYPASS"
GEMINI_OCR_MODEL = "gemini-2.5-flash"
GEMINI_OCR_MAX_RETRIES = 3
GEMINI_OCR_RETRY_BASE_SECONDS = 1.5

_TESSERACT_AVAILABLE: bool | None = None
_TESSERACT_FALLBACK_WARNING_EMITTED = False
_GEMINI_OCR_CLIENT: genai.Client | None = None
_GEMINI_OCR_CLIENT_API_KEY: str | None = None
_TRACE_FILE_MARKER_RE: re.Pattern[str] = re.compile(r"\[\[FILE:\s*[^\]]+?\s*\]\]")
_TRACE_PAGE_MARKER_RE: re.Pattern[str] = re.compile(r"\[\[PAGE\s+\d+\]\]")


def _resolve_gemini_api_key() -> str | None:
    try:
        from app.core.config import settings

        return (
            settings.GEMINI_API_KEY
            or settings.GOOGLE_API_KEY
            or os.environ.get("GEMINI_API_KEY")
            or os.environ.get("GOOGLE_API_KEY")
        )
    except Exception:
        return os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")


def _is_tesseract_available() -> bool:
    global _TESSERACT_AVAILABLE

    if _TESSERACT_AVAILABLE is not None:
        return _TESSERACT_AVAILABLE

    try:
        pytesseract.get_tesseract_version()
        _TESSERACT_AVAILABLE = True
    except Exception:
        _TESSERACT_AVAILABLE = False

    return _TESSERACT_AVAILABLE


def _get_gemini_ocr_client(api_key: str) -> genai.Client:
    global _GEMINI_OCR_CLIENT, _GEMINI_OCR_CLIENT_API_KEY

    if _GEMINI_OCR_CLIENT is not None and _GEMINI_OCR_CLIENT_API_KEY == api_key:
        return _GEMINI_OCR_CLIENT

    _GEMINI_OCR_CLIENT = genai.Client(api_key=api_key)
    _GEMINI_OCR_CLIENT_API_KEY = api_key
    return _GEMINI_OCR_CLIENT


class ExtractedArchiveFile(TypedDict):
    filename: str
    file_bytes: bytes


def _is_url(value: str) -> bool:
    parsed = urlparse(value)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def _looks_like_zip(data: bytes) -> bool:
    return data.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08"))


def _looks_like_rar(data: bytes) -> bool:
    return data.startswith(b"Rar!\x1a\x07\x00") or data.startswith(b"Rar!\x1a\x07\x01\x00")


def _looks_like_pdf(data: bytes) -> bool:
    return data.lstrip().startswith(b"%PDF")


def _looks_like_docx(data: bytes) -> bool:
    if not _looks_like_zip(data):
        return False

    try:
        with ZipFile(io.BytesIO(data), "r") as zip_file:
            return "word/document.xml" in zip_file.namelist()
    except BadZipFile:
        return False
    except Exception:
        return False


def _safe_decode_text(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "cp1251", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="ignore")


def _normalize_content_type(content_type: str | None) -> str:
    return (content_type or "").split(";", 1)[0].strip().lower()


def _source_filename(file_path: str | Path | None) -> str:
    value = str(file_path or "").strip()
    if not value or value == "<bytes>":
        return "uploaded_document"
    if value.startswith("<") and value.endswith(">"):
        return value
    return Path(value).name


def _format_file_marker(filename: str) -> str:
    return f"[[FILE: {filename}]]"


def _format_page_marker(page_number: int) -> str:
    return f"[[PAGE {page_number}]]"


def _format_page_text(filename: str, page_number: int, text: str) -> str:
    normalized = text.strip()
    if not normalized:
        return ""
    return f"{_format_file_marker(filename)}\n{_format_page_marker(page_number)}\n{normalized}"


def _format_single_page_document(filename: str, text: str) -> str:
    return _format_page_text(filename, 1, text)


def _has_trace_markers(text: str) -> bool:
    """Return True when text already has parser-style file and page markers."""
    return bool(_TRACE_FILE_MARKER_RE.search(text) and _TRACE_PAGE_MARKER_RE.search(text))


def ensure_trace_markers(filename: str, text: str) -> str:
    """
    Repair legacy markerless parsed text at compile time.

    DOCX files do not have reliable rendered page numbers, so page 1 means
    document-level provenance for single-page marker repair.
    """
    normalized = text.strip()
    if not normalized:
        return ""
    if _has_trace_markers(normalized):
        return normalized
    return _format_single_page_document(filename, normalized)


def _ocr_fallback_allowed(*, file_path: str, file_bytes: bytes) -> bool:
    suffix = Path(file_path or "").suffix.lower()
    if suffix in OCR_SAFE_SUFFIXES:
        return True
    return not suffix and _looks_like_pdf(file_bytes)


def _parse_extracted_documents(files: list[ExtractedArchiveFile]) -> str:
    parsed_chunks: list[str] = []

    for item in files:
        filename = item["filename"]
        suffix = Path(filename).suffix.lower()

        try:
            if suffix == ".pdf":
                text = parse_pdf(item["file_bytes"], file_path=filename)
            elif suffix == ".docx":
                text = parse_docx(item["file_bytes"], file_path=filename)
            elif suffix == ".txt":
                text = _format_single_page_document(
                    filename,
                    _safe_decode_text(item["file_bytes"]),
                )
            else:
                continue
        except Exception as exc:
            logger.error("Failed to parse '%s': %s", filename, exc, exc_info=True)
            continue

        normalized = text.strip()
        if normalized:
            parsed_chunks.append(normalized)

    return "\n\n".join(parsed_chunks).strip()


def _is_allowed_archive_member(member_name: str) -> bool:
    return member_name.lower().endswith(tuple(ALLOWED_EXTENSIONS))


def _flatten_archive_member_name(member_name: str) -> str:
    normalized_name = member_name.replace("\\", "/")
    return Path(normalized_name).name


def _build_unique_flat_path(extract_root: Path, flat_name: str) -> Path:
    base_name = Path(flat_name).stem
    suffix = Path(flat_name).suffix
    target_path = extract_root / flat_name

    counter = 1
    while target_path.exists():
        target_path = extract_root / f"{base_name}_{counter}{suffix}"
        counter += 1

    return target_path


def _extract_zip_contents_from_path(archive_path: Path) -> list[ExtractedArchiveFile]:
    extracted: list[ExtractedArchiveFile] = []
    with tempfile.TemporaryDirectory(prefix="zip_extract_") as temp_dir:
        extract_root = Path(temp_dir)
        with ZipFile(archive_path, "r") as zip_file:
            for member in zip_file.infolist():
                if member.is_dir():
                    continue
                if not _is_allowed_archive_member(member.filename):
                    continue
                flat_name = _flatten_archive_member_name(member.filename)
                if not flat_name:
                    logger.warning("Skipping zip member with invalid name: %s", member.filename)
                    continue
                target_path = _build_unique_flat_path(extract_root, flat_name)
                try:
                    with zip_file.open(member, "r") as source, target_path.open("wb") as target:
                        shutil.copyfileobj(source, target)
                    extracted.append({"filename": target_path.name, "file_bytes": target_path.read_bytes()})
                except Exception as exc:
                    logger.error("Failed to extract zip member '%s': %s", member.filename, exc, exc_info=True)
                    continue

    return extracted


def _extract_rar_contents_from_path(archive_path: Path) -> list[ExtractedArchiveFile]:
    extracted: list[ExtractedArchiveFile] = []
    with tempfile.TemporaryDirectory(prefix="rar_extract_") as temp_dir:
        extract_root = Path(temp_dir)
        with rarfile.RarFile(archive_path) as rar_archive:
            for member in rar_archive.infolist():
                if member.isdir():
                    continue
                if not _is_allowed_archive_member(member.filename):
                    continue
                flat_name = _flatten_archive_member_name(member.filename)
                if not flat_name:
                    logger.warning("Skipping rar member with invalid name: %s", member.filename)
                    continue
                target_path = _build_unique_flat_path(extract_root, flat_name)
                try:
                    with rar_archive.open(member, "r") as source, target_path.open("wb") as target:
                        shutil.copyfileobj(source, target)
                    extracted.append({"filename": target_path.name, "file_bytes": target_path.read_bytes()})
                except Exception as exc:
                    logger.error("Failed to extract rar member '%s': %s", member.filename, exc, exc_info=True)
                    continue

    return extracted


def _extract_archive_contents_from_bytes(archive_bytes: bytes, suffix: str) -> list[ExtractedArchiveFile]:
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as temp_file:
            temp_file.write(archive_bytes)
            temp_path = Path(temp_file.name)
        if suffix == ".zip":
            return _extract_zip_contents_from_path(temp_path)
        return _extract_rar_contents_from_path(temp_path)
    finally:
        if temp_path and temp_path.exists():
            try:
                temp_path.unlink()
            except OSError:
                logger.warning("Failed to remove temporary archive file: %s", temp_path)


def extract_archive_contents(archive_source: bytes | str | Path | BinaryIO) -> list[ExtractedArchiveFile]:
    """
    Extract only whitelisted document types from ZIP/RAR archives.
    Output filenames are flattened to a single temporary directory.
    """
    try:
        if isinstance(archive_source, (str, Path)):
            archive_path = Path(archive_source)
            suffix = archive_path.suffix.lower()
            archive_bytes = archive_path.read_bytes()

            if suffix == ".zip" or _looks_like_zip(archive_bytes):
                return _extract_archive_contents_from_bytes(archive_bytes, ".zip")
            if suffix == ".rar" or _looks_like_rar(archive_bytes):
                return _extract_archive_contents_from_bytes(archive_bytes, ".rar")

            try:
                return _extract_archive_contents_from_bytes(archive_bytes, ".zip")
            except BadZipFile:
                return _extract_archive_contents_from_bytes(archive_bytes, ".rar")

        archive_bytes: bytes
        if isinstance(archive_source, (bytes, bytearray)):
            archive_bytes = bytes(archive_source)
        elif hasattr(archive_source, "read"):
            archive_bytes = archive_source.read()
        else:
            logger.error("Unsupported archive source type: %s", type(archive_source))
            return []

        if _looks_like_zip(archive_bytes):
            return _extract_archive_contents_from_bytes(archive_bytes, ".zip")
        if _looks_like_rar(archive_bytes):
            return _extract_archive_contents_from_bytes(archive_bytes, ".rar")

        try:
            return _extract_archive_contents_from_bytes(archive_bytes, ".zip")
        except BadZipFile:
            return _extract_archive_contents_from_bytes(archive_bytes, ".rar")
    except (BadZipFile, rarfile.Error, FileNotFoundError, OSError) as exc:
        logger.error("Archive extraction failed: %s", exc, exc_info=True)
        return []
    except Exception as exc:
        logger.error("Unexpected archive extraction error: %s", exc, exc_info=True)
        return []


def parse_docx(docx_bytes: bytes, file_path: str | Path = "<bytes>") -> str:
    """Parse DOCX bytes and return full text."""
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        logger.error("DOCX open failed: %s", exc, exc_info=True)
        return ""

    chunks: list[str] = []

    try:
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        chunks.append(text)
    except Exception as exc:
        logger.error("DOCX parsing failed: %s", exc, exc_info=True)
        return ""

    return _format_single_page_document(
        _source_filename(file_path),
        "\n".join(chunks).strip(),
    )


def _page_has_visual_content(page: fitz.Page) -> bool:
    try:
        if page.get_images(full=True):
            return True
    except Exception:
        pass

    try:
        if page.get_drawings():
            return True
    except Exception:
        pass

    return False


def _ocr_pdf_page(image: Image.Image) -> str:
    global _TESSERACT_AVAILABLE, _TESSERACT_FALLBACK_WARNING_EMITTED

    if _is_tesseract_available():
        try:
            return pytesseract.image_to_string(
                image,
                lang=OCR_LANGS,
                timeout=OCR_PAGE_TIMEOUT_SECONDS,
            ).strip()
        except Exception:
            _TESSERACT_AVAILABLE = False

    if not _TESSERACT_FALLBACK_WARNING_EMITTED:
        logger.warning("Tesseract not found. Falling back to Gemini Vision OCR for page.")
        _TESSERACT_FALLBACK_WARNING_EMITTED = True

    api_key = _resolve_gemini_api_key()
    if not api_key:
        logger.error("Gemini Vision OCR fallback is unavailable: GEMINI_API_KEY is not set.")
        return ""

    prompt = (
        "Extract all text and tables from this document page precisely as they appear. "
        "Return strictly the extracted text and nothing else."
    )
    client = _get_gemini_ocr_client(api_key)

    for attempt in range(1, GEMINI_OCR_MAX_RETRIES + 1):
        try:
            response = client.models.generate_content(
                model=GEMINI_OCR_MODEL,
                contents=[prompt, image],
            )
            return (getattr(response, "text", "") or "").strip()
        except Exception as exc:
            if attempt >= GEMINI_OCR_MAX_RETRIES:
                logger.error(
                    "Gemini Vision OCR fallback failed after %s attempts: %s",
                    GEMINI_OCR_MAX_RETRIES,
                    exc,
                    exc_info=True,
                )
                return ""

            delay_seconds = GEMINI_OCR_RETRY_BASE_SECONDS * attempt
            logger.warning(
                "Gemini Vision OCR transient failure, retrying attempt %s/%s in %.1fs: %s",
                attempt + 1,
                GEMINI_OCR_MAX_RETRIES,
                delay_seconds,
                exc,
            )
            time.sleep(delay_seconds)

    return ""


def _ocr_pdf_page_from_pdf_bytes(pdf_bytes: bytes, page_number: int) -> str:
    import concurrent.futures

    def _render_and_ocr() -> str:
        with tempfile.TemporaryDirectory(prefix="pdf_ocr_page_") as temp_dir:
            image_paths = convert_from_bytes(
                pdf_bytes,
                dpi=OCR_RENDER_DPI,
                first_page=page_number,
                last_page=page_number,
                fmt="png",
                output_folder=temp_dir,
                paths_only=True,
                thread_count=1,
                timeout=OCR_PAGE_TIMEOUT_SECONDS,
            )
            if not image_paths:
                return ""

            page_chunks: list[str] = []
            for image_path in image_paths:
                try:
                    with Image.open(image_path) as image:
                        text = _ocr_pdf_page(image)
                    if text:
                        page_chunks.append(text)
                except Exception as exc:
                    logger.error("OCR image read failed (page %s): %s", page_number, exc, exc_info=True)
                    continue

            return "\n".join(page_chunks).strip()

    try:
        executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
        future = executor.submit(_render_and_ocr)
        try:
            return future.result(timeout=OCR_PAGE_TIMEOUT_SECONDS)
        finally:
            executor.shutdown(wait=future.done(), cancel_futures=True)
    except concurrent.futures.TimeoutError:
        logger.warning(
            "[SONAR] OCR timed out after %ss on page %s — skipping.",
            OCR_PAGE_TIMEOUT_SECONDS,
            page_number,
        )
        return ""
    except Exception as exc:
        logger.error("OCR conversion failed (page %s): %s", page_number, exc, exc_info=True)
        return ""


def _ocr_entire_pdf(pdf_bytes: bytes, file_path: str = "<bytes>") -> str:
    if not OCR_ENABLED:
        logger.warning("[SONAR] Full-document OCR skipped because TENDER_OCR_DISABLED is active.")
        return ""

    try:
        with tempfile.TemporaryDirectory(prefix="pdf_ocr_full_") as temp_dir:
            image_paths = convert_from_bytes(
                pdf_bytes,
                dpi=OCR_RENDER_DPI,
                fmt="png",
                output_folder=temp_dir,
                paths_only=True,
                thread_count=1,
                timeout=OCR_PAGE_TIMEOUT_SECONDS,
            )
            if not image_paths:
                return ""

            filename = _source_filename(file_path)
            page_texts: list[str] = []
            for index, image_path in enumerate(image_paths, start=1):
                try:
                    with Image.open(image_path) as image:
                        text = _ocr_pdf_page(image)
                    formatted = _format_page_text(filename, index, text)
                    if formatted:
                        page_texts.append(formatted)
                except Exception as exc:
                    logger.error("OCR failed (fallback page %s): %s", index, exc, exc_info=True)
                    continue

            return "\n\n".join(page_texts).strip()
    except Exception as exc:
        logger.error("Full-document OCR failed: %s", exc, exc_info=True)
        return ""


def parse_pdf(pdf_bytes: bytes, file_path: str = "<bytes>") -> str:
    """
    Parse PDF bytes using native extraction and OCR fallback for scanned pages.
    """
    if _looks_like_zip(pdf_bytes) or _looks_like_rar(pdf_bytes):
        logger.warning(
            "[SONAR] Refusing to send archive bytes to PyMuPDF: %s",
            file_path,
        )
        return ""

    filename = _source_filename(file_path)
    page_texts: list[str] = []
    ocr_fallback_logged = False
    ocr_disabled_logged = False
    ocr_text_threshold_logged = False
    ocr_pages_processed = 0
    extracted_chars_so_far = 0
    demo_ocr_bypass = _env_flag(DEMO_OCR_BYPASS_ENV)
    if demo_ocr_bypass:
        logger.warning(
            "[SONAR] %s=True; OCR page cap bypass is active for %s.",
            DEMO_OCR_BYPASS_ENV,
            filename,
        )

    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_doc:
            logger.info("[SONAR] PyMuPDF opened the document successfully.")
            for page_index, page in enumerate(pdf_doc, start=1):
                native_text = ""
                try:
                    native_text = (page.get_text() or "").strip()
                except Exception as exc:
                    logger.error(
                        "Native PDF extraction failed (page %s): %s",
                        page_index,
                        exc,
                        exc_info=True,
                    )

                if len(native_text) >= OCR_MIN_TEXT_LEN and not demo_ocr_bypass:
                    page_texts.append(_format_page_text(filename, page_index, native_text))
                    extracted_chars_so_far += len(native_text)
                    continue

                ocr_text = ""
                if demo_ocr_bypass or native_text or _page_has_visual_content(page):
                    run_ocr = True
                    if not demo_ocr_bypass and not OCR_ENABLED:
                        if not ocr_disabled_logged:
                            logger.info("[SONAR] OCR disabled; parsing native PDF text only.")
                            ocr_disabled_logged = True
                        run_ocr = False
                    elif (
                        not demo_ocr_bypass
                        and OCR_SKIP_AFTER_TEXT_CHARS
                        and extracted_chars_so_far >= OCR_SKIP_AFTER_TEXT_CHARS
                    ):
                        if not ocr_text_threshold_logged:
                            logger.info(
                                "[SONAR] OCR skipped after %s native characters at page %s.",
                                extracted_chars_so_far,
                                page_index,
                            )
                            ocr_text_threshold_logged = True
                        run_ocr = False
                    elif not demo_ocr_bypass and ocr_pages_processed >= OCR_MAX_PAGES:
                        logger.warning(
                            "[SONAR] OCR page limit (%s) reached at page %s — skipping remaining pages.",
                            OCR_MAX_PAGES,
                            page_index,
                        )
                        continue
                    if run_ocr:
                        if not ocr_fallback_logged:
                            logger.info("[SONAR] Empty text detected. Triggering OCR fallback.")
                            ocr_fallback_logged = True
                        ocr_text = _ocr_pdf_page_from_pdf_bytes(pdf_bytes, page_index)
                        ocr_pages_processed += 1

                merged = "\n".join(part for part in (native_text, ocr_text) if part).strip()
                formatted = _format_page_text(filename, page_index, merged)
                if formatted:
                    page_texts.append(formatted)
                    extracted_chars_so_far += len(merged)
    except Exception as e:
        logger.error(f"[SONAR] EXTRACTION FAILED on {file_path}. Reason: {str(e)}", exc_info=True)
        if _ocr_fallback_allowed(file_path=file_path, file_bytes=pdf_bytes):
            return _ocr_entire_pdf(pdf_bytes, file_path=file_path)
        logger.warning("[SONAR] OCR fallback skipped for unsupported non-PDF source: %s", file_path)
        return ""

    extracted_text = "\n\n".join(page_texts).strip()
    logger.info(f"[SONAR] Extraction complete. Total characters: {len(extracted_text)}")
    return extracted_text


async def process_tender_document(
    source: bytes | str | Path,
    filename: str | None = None,
) -> str:
    """
    Async entry point:
    - accepts URL, local file path, or raw bytes
    - routes to archive extraction or direct parser
    - returns concatenated legally relevant text
    """
    payload: bytes = b""
    content_type = ""
    inferred_name = filename or ""
    path_source: Path | None = None

    try:
        if isinstance(source, (bytes, bytearray)):
            payload = bytes(source)
        elif isinstance(source, Path):
            path_source = source
            inferred_name = inferred_name or source.name
            payload = source.read_bytes()
        elif isinstance(source, str):
            if _is_url(source):
                inferred_name = inferred_name or Path(urlparse(source).path).name
                async with httpx.AsyncClient(timeout=60.0) as client:
                    response = await client.get(source)
                    response.raise_for_status()
                    content_type = _normalize_content_type(response.headers.get("content-type"))
                    payload = response.content
            else:
                path_source = Path(source)
                inferred_name = inferred_name or path_source.name
                payload = path_source.read_bytes()
        else:
            logger.error("Unsupported source type for parser: %s", type(source))
            return ""
    except (httpx.HTTPError, FileNotFoundError, OSError) as exc:
        logger.error("Failed to load source document: %s", exc, exc_info=True)
        return ""
    except Exception as exc:
        logger.error("Unexpected source loading error: %s", exc, exc_info=True)
        return ""

    if not payload:
        return ""

    # Route container formats from their binary signatures before any
    # filename- or MIME-based heuristics. This prevents PyMuPDF from
    # opening RAR payloads as CBR and triggering OCR on compressed data.
    is_rar = _looks_like_rar(payload)
    is_docx = _looks_like_docx(payload)
    is_zip_archive = _looks_like_zip(payload) and not is_docx

    if is_zip_archive or is_rar:
        archive_input: bytes | str | Path = path_source if path_source else payload
        extracted_files = await asyncio.to_thread(extract_archive_contents, archive_input)
        return await asyncio.to_thread(_parse_extracted_documents, extracted_files)

    suffix = Path(inferred_name).suffix.lower()
    normalized_content_type = _normalize_content_type(content_type)

    if suffix in ARCHIVE_SUFFIXES or normalized_content_type in ARCHIVE_MIME_TYPES:
        archive_input: bytes | str | Path = path_source if path_source else payload
        extracted_files = await asyncio.to_thread(extract_archive_contents, archive_input)
        return await asyncio.to_thread(_parse_extracted_documents, extracted_files)

    if suffix == ".docx" or normalized_content_type in DOCX_MIME_TYPES or is_docx:
        source_label = inferred_name or str(path_source) if path_source else inferred_name or "<bytes>"
        return await asyncio.to_thread(parse_docx, payload, source_label)

    if suffix == ".txt" or normalized_content_type.startswith("text/"):
        return _format_single_page_document(
            _source_filename(inferred_name or path_source or "<bytes>"),
            _safe_decode_text(payload),
        )

    if suffix == ".pdf" or normalized_content_type in PDF_MIME_TYPES or _looks_like_pdf(payload):
        source_label = str(path_source) if path_source else inferred_name or "<bytes>"
        return await asyncio.to_thread(parse_pdf, payload, source_label)

    logger.warning(
        "Unsupported source format for process_tender_document: name=%s content_type=%s",
        inferred_name or "<bytes>",
        normalized_content_type or "<unknown>",
    )
    return ""


def extract_text_from_pdf(file_path: str | Path) -> str:
    """
    Backward-compatible wrapper for existing code paths.
    """
    try:
        path = Path(file_path)
        logger.info(f"[SONAR] Attempting to extract: {file_path} | Size: {os.path.getsize(path)} bytes")
        return parse_pdf(path.read_bytes(), file_path=str(path))
    except Exception as e:
        logger.error(f"[SONAR] EXTRACTION FAILED on {file_path}. Reason: {str(e)}", exc_info=True)
        return ""


def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
    """
    Backward-compatible wrapper for existing code paths.
    """
    normalized = file_type.lower().strip(".")
    try:
        if normalized == "pdf":
            return parse_pdf(file_bytes, file_path="uploaded_document.pdf")
        if normalized in {"doc", "docx"}:
            return parse_docx(file_bytes, file_path=f"uploaded_document.{normalized}")
        if normalized == "txt":
            return _format_single_page_document("uploaded_document.txt", _safe_decode_text(file_bytes))
        if normalized in {"zip", "rar"}:
            extracted = extract_archive_contents(file_bytes)
            return _parse_extracted_documents(extracted)
    except Exception as exc:
        logger.error("extract_text_from_bytes failed for type '%s': %s", normalized, exc, exc_info=True)
        return ""

    logger.warning("Unsupported file type in extract_text_from_bytes: %s", file_type)
    return ""


def extract_text_from_file(file_path: str | Path) -> str:
    """
    Backward-compatible wrapper for existing code paths.
    """
    path = Path(file_path)
    suffix = path.suffix.lower().strip(".")
    try:
        file_bytes = path.read_bytes()
    except (FileNotFoundError, OSError) as exc:
        logger.error("Failed to read file '%s': %s", file_path, exc, exc_info=True)
        return ""
    except Exception as exc:
        logger.error("Unexpected file read error '%s': %s", file_path, exc, exc_info=True)
        return ""

    if suffix == "pdf":
        return parse_pdf(file_bytes, file_path=str(path))
    if suffix in {"doc", "docx"}:
        return parse_docx(file_bytes, file_path=str(path))
    if suffix == "txt":
        return _format_single_page_document(path.name, _safe_decode_text(file_bytes))
    if suffix in {"zip", "rar"}:
        extracted = extract_archive_contents(file_bytes)
        return _parse_extracted_documents(extracted)

    return extract_text_from_bytes(file_bytes, suffix)
